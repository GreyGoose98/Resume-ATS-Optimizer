import inspect
if not hasattr(inspect, 'getargspec'):
    inspect.getargspec = inspect.getfullargspec

import streamlit as st
import tempfile
import os
import base64
import re
import requests

# Import google generative AI (Gemini)
import google.generativeai as genai

# Define a custom GeminiModel to wrap the Gemini API
class GeminiModel:
    def __init__(self, id):
        self.id = id
        # Uncomment and update the API key if required:
        # genai.configure(api_key="YOUR_GEMINI_API_KEY")
    def generate_content(self, prompt):
        model = genai.GenerativeModel(self.id)
        response = model.generate_content(prompt)
        return response

# Attempt to import phi.agent; if not found, define a dummy Agent class.
try:
    from phi.agent import Agent
except ModuleNotFoundError:
    class Agent:
        def __init__(self, name, model, instructions, show_tool_calls, markdown):
            self.name = name
            self.model = model
            self.instructions = instructions
            self.show_tool_calls = show_tool_calls
            self.markdown = markdown
        def print_response(self, prompt_text):
            response = self.model.generate_content(prompt_text)
            return response.text

# --- Define Agents using GeminiModel ---
analysis_agent = Agent(
    name="Analysis Agent",
    model=GeminiModel(id="gemini-1.5-flash"),
    instructions=["Provide a detailed ATS analysis comparing the resume with the job description using the checklist."],
    show_tool_calls=True,
    markdown=True,
)

boost_agent = Agent(
    name="Boost Agent",
    model=GeminiModel(id="gemini-1.5-flash"),
    instructions=["Revise the resume to improve its ATS compatibility based on the provided analysis report. Preserve details and improve formatting."],
    show_tool_calls=True,
    markdown=True,
)

custom_agent = Agent(
    name="Custom Update Agent",
    model=GeminiModel(id="gemini-1.5-flash"),
    instructions=["Update the resume strictly following the custom instructions provided. Ensure professional tone and formatting."],
    show_tool_calls=True,
    markdown=True,
)

create_agent = Agent(
    name="Create Resume Agent",
    model=GeminiModel(id="gemini-1.5-flash"),
    instructions=["Generate a professional resume in Markdown format using the provided information."],
    show_tool_calls=True,
    markdown=True,
)

def call_agent(prompt_text, agent):
    # Helper function to call a specified agent and return its response text.
    return agent.print_response(prompt_text)

# --- Helper Functions ---
def clean_placeholder_text(text):
    """
    Remove common placeholder phrases from the text.
    """
    placeholders = [
        r"add relevant experience",
        r"add your experience here",
        r"placeholder",
    ]
    for ph in placeholders:
        text = re.sub(ph, "", text, flags=re.IGNORECASE)
    return text

def generate_docx_from_markdown(markdown_text):
    """
    Convert Markdown text to a DOCX binary using python-docx.
    A simple parser is used to create headings, bullet lists, and paragraphs.
    """
    from docx import Document
    document = Document()
    lines = markdown_text.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            document.add_paragraph("")
        elif line.startswith("#"):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            if level == 1:
                document.add_heading(heading_text, level=1)
            elif level == 2:
                document.add_heading(heading_text, level=2)
            else:
                document.add_heading(heading_text, level=3)
        elif line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            document.add_paragraph(line)
    f = BytesIO()
    document.save(f)
    f.seek(0)
    return f.read()

def analyze_resume(resume_text, jd_text):
    prompt = (f"""Analyze the following resume with respect to the job description below.
Use the following checklist for guidance:
{check_list}
Provide the ATS score (0 to 100) as a floating point number with a breakdown of scores per section, and a detailed improvement report.
Return the output in the following format:

ATS Score : <ATS SCORE>
Detailed Report: <DETAILED REPORT>

Resume:
{resume_text}

Job Description:
{jd_text}""")
    analysis = call_agent(prompt, analysis_agent)
    analysis = clean_placeholder_text(analysis)
    return analysis

def boost_resume_md(resume_text, jd_text, analysis_report):
    prompt = (
        f"""You are a highly skillful tool that boosts and enhances resumes by integrating recommendations from an analysis report.
Revise the resume to improve its ATS score, compatibility, and formatting while preserving its details and style.
Return the updated resume in Markdown format with proper headings, bullet points, and styling.

ATS Analysis Report:
{analysis_report}

Resume:
{resume_text}

Job Description:
{jd_text}

Return only the updated resume in Markdown format.
Strict guideline: Return only the updated resume text in a professional tone."""
    )
    boosted_md = call_agent(prompt, boost_agent)
    boosted_md = clean_placeholder_text(boosted_md)
    return boosted_md

def custom_update_resume(resume_text, custom_prompt):
    prompt = (
        f"""You are a professional resume editor. 
Using the following resume, update it strictly according to the custom instructions provided.
Resume:
{resume_text}

Custom Instructions:
{custom_prompt}

Return only the updated resume in Markdown format in a professional tone."""
    )
    updated_resume = call_agent(prompt, custom_agent)
    updated_resume = clean_placeholder_text(updated_resume)
    return updated_resume

def create_resume_from_form(form_data):
    prompt = (
        f"""Based on the following information, create a professional resume in Markdown format.
Use clear headings, bullet points, and a professional tone. Include all mandatory details and incorporate optional sections as provided.

Name: {form_data.get('name', '')}
Email: {form_data.get('email', '')}
Phone: {form_data.get('phone', '')}
LinkedIn: {form_data.get('linkedin', '')}
Address: {form_data.get('address', '')}

Education:
{form_data.get('education', '')}

Work Experience:
{form_data.get('experience', '')}

Skills:
{form_data.get('skills', '')}

Projects (Optional):
{form_data.get('projects', '')}

Certifications (Optional):
{form_data.get('certifications', '')}

Achievements (Optional):
{form_data.get('achievements', '')}

Hobbies (Optional):
{form_data.get('hobbies', '')}

Return only the resume in Markdown format."""
    )
    new_resume = call_agent(prompt, create_agent)
    new_resume = clean_placeholder_text(new_resume)
    return new_resume

def extract_text_from_pdf(file_bytes):
    text = ""
    try:
        if pdfminer_extract_text:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_filename = tmp.name
            text = pdfminer_extract_text(tmp_filename)
            os.remove(tmp_filename)
        else:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file_bytes):
    text = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_filename = tmp.name
        # Check if 'mammoth' exists in globals() and use it if available.
        if 'mammoth' in globals() and mammoth is not None:
            with open(tmp_filename, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                text = result.value
        else:
            doc = docx.Document(tmp_filename)
            for para in doc.paragraphs:
                text += para.text + "\n"
        os.remove(tmp_filename)
    except Exception as e:
        st.error(f"Error extracting formatted text from DOCX: {e}")
    return text

st.set_page_config(
    page_title="Resume ATS Optimizer Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main {
        padding: 2rem 3rem;
        background-color: #f8f9fa;
    }
    .stButton > button {
        width: 100%;
        height: 3rem;
        font-weight: 600;
        border-radius: 8px;
        margin-top: 1rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    .stDownloadButton > button {
        width: 100%;
        border-radius: 8px;
        margin-top: 0.5rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .results-container {
        max-height: 500px;
        overflow-y: auto;
        padding: 1rem;
        background-color: #f8f9fa;
        border-radius: 8px;
        border: 1px solid #e9ecef;
    }
    .logo-text {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(90deg, #1E88E5, #4CAF50);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 2rem;
    }
    .score-container {
        display: flex;
        justify-content: center;
        align-items: center;
        margin: 2rem 0;
    }
    .score-circle {
        width: 150px;
        height: 150px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        background: conic-gradient(#4CAF50 var(--percentage), #f3f3f3 var(--percentage));
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
    }
    .score-inner {
        width: 130px;
        height: 130px;
        border-radius: 50%;
        background: white;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 2.5rem;
        font-weight: 700;
        color: #4CAF50;
    }
    .section-title {
        font-weight: 600;
        margin-bottom: 1.5rem;
        color: #1E88E5;
        border-left: 4px solid #1E88E5;
        padding-left: 10px;
    }
</style>
""", unsafe_allow_html=True)

check_list = """Checklist for ATS Score Calculation and Improvement  

1. Resume Parsing  
   - Extract text content and preserve formatting (headings, bullet points, etc.).  
   - Identify key sections: Work Experience, Education, Skills, Certifications.  

2. Job Description Analysis  
   - Extract required keywords and phrases from the job description.  
   - Identify critical skills, certifications, and experience requirements.  
   - Highlight industry-specific terminology.  

3. Keyword Matching  
   - Compare keywords from the job description to those found in the resume.  
   - Count the frequency of keyword occurrences.  
   - Ensure keywords are detected in context (e.g., within relevant sections).  

4. Contextual and Relevance Evaluation  
   - Assess if keywords appear in appropriate sections.  
   - Evaluate context for meaningful usage rather than repetition.  

5. Formatting and Structure Assessment  
   - Check for ATS-friendly formatting (clear headings, bullet points, simple layout).  
   - Verify that the resume avoids excessive graphics, tables, or unusual elements.  

6. Scoring Algorithm  
   - Define weighting for each component (keyword match, context, structure).  
   - Aggregate scores to generate an overall ATS score with a detailed breakdown.  
   - Provide a breakdown of scores for different sections.  

7. Detailed Improvement Report  
   - Highlight missing or underrepresented keywords.  
   - Provide recommendations to enhance keyword placement and formatting.  

8. Error Handling and Validation  
   - Ensure accurate text extraction without loss of key formatting.  
   - Handle parsing errors gracefully with user-friendly messages.  

9. Testing and Refinement  
   - Test the application with diverse resume formats and job descriptions.  
   - Regularly update the application to handle new resume trends and ATS criteria.
"""

def main():
    if 'analysis_report' not in st.session_state:
        st.session_state.analysis_report = None
    if 'boosted_resume' not in st.session_state:
        st.session_state.boosted_resume = None
    if 'ats_score' not in st.session_state:
        st.session_state.ats_score = 0.0
    if 'boosted_ats_score' not in st.session_state:
        st.session_state.boosted_ats_score = 0.0
    if 'custom_updated_resume' not in st.session_state:
        st.session_state.custom_updated_resume = None
    if 'new_resume' not in st.session_state:
        st.session_state.new_resume = None

    st.markdown('<div class="logo-text">Resume ATS Optimizer Pro</div>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; margin-bottom: 3rem;">Boost your resume\'s chances of getting past Applicant Tracking Systems</p>', unsafe_allow_html=True)
    
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Upload & Analyze", "✏️ Custom Update", "🆕 Create New Resume", "📊 Results & Downloads"])
    
    with tab1:
        resume_text = ""
        jd_text = ""
        
        col1, col2 = st.columns(2)
        
        with col1:
            with st.container():
                st.markdown('<h3 class="section-title">Upload Your Resume</h3>', unsafe_allow_html=True)
                resume_file = st.file_uploader("Choose your resume file", type=["pdf", "tex", "txt", "docx"], key="resume_uploader")
                if resume_file is not None:
                    st.success(f"✅ Uploaded: {resume_file.name}")
                    file_details = resume_file.name.lower()
                    file_content = resume_file.read()
                    try:
                        if file_details.endswith(".tex") or file_details.endswith(".txt"):
                            resume_text = file_content.decode("utf-8")
                        elif file_details.endswith(".pdf"):
                            resume_text = extract_text_from_pdf(file_content)
                        elif file_details.endswith(".docx"):
                            resume_text = extract_text_from_docx(file_content)
                        else:
                            st.warning("Unsupported file type for resume")
                    except Exception as e:
                        st.error(f"Error processing file: {e}")
                else:
                    st.info("Please upload your resume file (PDF, TEX, TXT, or DOCX)")
                
                if resume_text:
                    if st.checkbox("Display extracted resume content", key="show_resume"):
                        with st.expander("Extracted Resume Content", expanded=True):
                            st.text_area("Resume Content", resume_text, height=200, disabled=True)
        
        with col2:
            with st.container():
                st.markdown('<h3 class="section-title">Upload or Enter Job Description</h3>', unsafe_allow_html=True)
                jd_file = st.file_uploader("Choose job description file", type=["txt", "pdf", "docx"], key="jd_uploader")
                if jd_file is not None:
                    st.success(f"✅ Uploaded: {jd_file.name}")
                    jd_details = jd_file.name.lower()
                    jd_file_content = jd_file.read()
                    try:
                        if jd_details.endswith(".txt"):
                            jd_text = jd_file_content.decode("utf-8")
                        elif jd_details.endswith(".pdf"):
                            jd_text = extract_text_from_pdf(jd_file_content)
                        elif jd_details.endswith(".docx"):
                            jd_text = extract_text_from_docx(jd_file_content)
                        else:
                            st.warning("Unsupported file type for job description")
                    except Exception as e:
                        st.error(f"Error processing file: {e}")
                jd_text_input = st.text_area("Or, paste your Job Description text here", height=150)
                if jd_text_input.strip():
                    jd_text = jd_text_input
                
                if jd_text:
                    if st.checkbox("Display extracted job description content", key="show_jd"):
                        with st.expander("Extracted Job Description Content", expanded=True):
                            st.text_area("Job Description Content", jd_text, height=200, disabled=True)
                if not jd_text:
                    st.info("Please provide the job description either via file upload or text input.")
        
        with st.container():
            st.markdown('<h3 class="section-title">Analyze Your Resume</h3>', unsafe_allow_html=True)
            analyze_col, boost_col = st.columns(2)
            with analyze_col:
                analyze_button = st.button("📊 Analyze Resume", use_container_width=True, disabled=not (resume_text and jd_text))
            with boost_col:
                boost_button = st.button("🚀 Boost Resume", use_container_width=True, disabled=(st.session_state.analysis_report is None))
            if analyze_button:
                st.session_state.analysis_report = analyze_resume(resume_text, jd_text)
                st.session_state.ats_score = extract_ats_score(st.session_state.analysis_report)
            if boost_button:
                st.session_state.boosted_resume = boost_resume_md(resume_text, jd_text, st.session_state.analysis_report)
                boosted_analysis = analyze_resume(st.session_state.boosted_resume, jd_text)
                st.session_state.boosted_ats_score = extract_ats_score(boosted_analysis)
    
    with tab2:
        st.markdown('<h3 class="section-title">Custom Update Your Resume</h3>', unsafe_allow_html=True)
        custom_resume_text = ""
        uploaded_file = st.file_uploader("Upload your resume for custom update", type=["pdf", "tex", "txt", "docx"], key="custom_resume")
        if uploaded_file is not None:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            details = uploaded_file.name.lower()
            content = uploaded_file.read()
            try:
                if details.endswith(".tex") or details.endswith(".txt"):
                    custom_resume_text = content.decode("utf-8")
                elif details.endswith(".pdf"):
                    custom_resume_text = extract_text_from_pdf(content)
                elif details.endswith(".docx"):
                    custom_resume_text = extract_text_from_docx(content)
                else:
                    st.warning("Unsupported file type")
            except Exception as e:
                st.error(f"Error processing file: {e}")
        else:
            st.info("Please upload your resume for custom update")
        
        custom_prompt = st.text_area("Enter custom update instructions", 
                                     "E.g., update the skills section, emphasize recent projects, and improve formatting.",
                                     height=150)
        if st.button("Apply Custom Update"):
            if custom_resume_text and custom_prompt.strip():
                st.session_state.custom_updated_resume = custom_update_resume(custom_resume_text, custom_prompt)
            else:
                st.error("Please upload a resume and enter update instructions.")
        
        if st.session_state.get("custom_updated_resume"):
            st.markdown("### Custom Updated Resume")
            st.markdown(st.session_state.custom_updated_resume)
    
    with tab3:
        st.markdown('<h3 class="section-title">Create a New Resume from Scratch</h3>', unsafe_allow_html=True)
        st.info("Please fill out the form below. Mandatory fields are marked with *.")
        with st.form("new_resume_form"):
            name = st.text_input("Full Name *")
            email = st.text_input("Email *")
            phone = st.text_input("Phone Number *")
            linkedin = st.text_input("LinkedIn URL (Optional)")
            address = st.text_input("Address (Optional)")
            education = st.text_area("Education (e.g., degrees, institutions, dates) *")
            experience = st.text_area("Work Experience (e.g., roles, companies, dates) *")
            skills = st.text_area("Skills *")
            projects = st.text_area("Projects (Optional)")
            certifications = st.text_area("Certifications (Optional)")
            achievements = st.text_area("Achievements (Optional)")
            hobbies = st.text_input("Hobbies (Optional)")
            submit_button = st.form_submit_button("Create Resume")
            
            if submit_button:
                if not name or not email or not phone or not education or not experience or not skills:
                    st.error("Please fill in all mandatory fields marked with *.")
                else:
                    form_data = {
                        "name": name,
                        "email": email,
                        "phone": phone,
                        "linkedin": linkedin,
                        "address": address,
                        "education": education,
                        "experience": experience,
                        "skills": skills,
                        "projects": projects,
                        "certifications": certifications,
                        "achievements": achievements,
                        "hobbies": hobbies,
                    }
                    st.session_state.new_resume = create_resume_from_form(form_data)
        
        if st.session_state.get("new_resume"):
            st.markdown("### Newly Created Resume")
            st.markdown(st.session_state.new_resume)
    
    with tab4:
        if st.session_state.get('analysis_report') is not None:
            with st.container():
                st.markdown('<h3 class="section-title">Analysis Results</h3>', unsafe_allow_html=True)
                score = st.session_state.ats_score
                st.markdown(
                    f'''
                    <div class="score-container">
                        <div class="score-circle" style="--percentage: {score}%">
                            <div class="score-inner">{score}</div>
                        </div>
                    </div>
                    <p style="text-align: center; font-weight: 600; margin-bottom: 2rem;">Original ATS Compatibility Score</p>
                    ''', 
                    unsafe_allow_html=True
                )
                with st.expander("📋 View Detailed Analysis Report", expanded=True):
                    st.markdown('<div class="results-container">', unsafe_allow_html=True)
                    formatted_report = st.session_state.analysis_report.replace("ATS Score :", "").replace("Detailed Report:", "**Detailed Report:**")
                    st.markdown(formatted_report)
                    st.markdown('</div>', unsafe_allow_html=True)
            
            if st.session_state.get('boosted_resume') is not None:
                with st.container():
                    st.markdown('<h3 class="section-title">Optimized Resume</h3>', unsafe_allow_html=True)
                    with st.expander("📝 View Optimized Resume", expanded=True):
                        st.markdown('<div class="results-container">', unsafe_allow_html=True)
                        st.markdown(st.session_state.boosted_resume)
                        st.markdown('</div>', unsafe_allow_html=True)
                with st.container():
                    boosted_score = st.session_state.boosted_ats_score
                    st.markdown(
                        f'''
                        <div class="score-container">
                            <div class="score-circle" style="--percentage: {boosted_score}%">
                                <div class="score-inner">{boosted_score}</div>
                            </div>
                        </div>
                        <p style="text-align: center; font-weight: 600; margin-bottom: 2rem;">Optimized ATS Compatibility Score</p>
                        ''', 
                        unsafe_allow_html=True
                    )
            
            if st.session_state.get("custom_updated_resume"):
                with st.container():
                    st.markdown('<h3 class="section-title">Custom Updated Resume</h3>', unsafe_allow_html=True)
                    with st.expander("📝 View Custom Updated Resume", expanded=True):
                        st.markdown('<div class="results-container">', unsafe_allow_html=True)
                        st.markdown(st.session_state.custom_updated_resume)
                        st.markdown('</div>', unsafe_allow_html=True)
            
            if st.session_state.get("new_resume"):
                with st.container():
                    st.markdown('<h3 class="section-title">Newly Created Resume</h3>', unsafe_allow_html=True)
                    with st.expander("📝 View Newly Created Resume", expanded=True):
                        st.markdown('<div class="results-container">', unsafe_allow_html=True)
                        st.markdown(st.session_state.new_resume)
                        st.markdown('</div>', unsafe_allow_html=True)
            
            with st.container():
                st.markdown('<h3 class="section-title">Download Options</h3>', unsafe_allow_html=True)
                col1, col2 = st.columns(2)
                with col1:
                    if st.session_state.get("boosted_resume"):
                        boosted_html = markdown.markdown(st.session_state.boosted_resume)
                        st.download_button(
                            "📥 Download Optimized Resume as Markdown", 
                            data=st.session_state.boosted_resume, 
                            file_name="optimized_resume.md", 
                            mime="text/markdown",
                            use_container_width=True
                        )
                        st.download_button(
                            "📥 Download Optimized Resume as HTML", 
                            data=boosted_html, 
                            file_name="optimized_resume.html", 
                            mime="text/html",
                            use_container_width=True
                        )
                with col2:
                    if st.session_state.get("boosted_resume"):
                        try:
                            docx_resume = generate_docx_from_markdown(st.session_state.boosted_resume)
                            st.download_button(
                                "📥 Download Optimized Resume as Word Document", 
                                data=docx_resume, 
                                file_name="optimized_resume.docx", 
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        except Exception as e:
                            st.error(f"Conversion to Word failed: {e}")
                    st.info("To save as PDF: Use your browser's print functionality and select 'Save as PDF'")
        else:
            st.info("No analysis results yet. Please go to 'Upload & Analyze' tab and analyze your resume first.")

    def load_image_as_base64(file_path: str) -> str:
        try:
            with open(file_path, "rb") as file:
                data = file.read()
            return base64.b64encode(data).decode("utf-8")
        except Exception as e:
            st.error(f"Error loading image from {file_path}: {e}")
            return ""

    logo_path = "inventify_logo.png"
    if os.path.exists(logo_path):
        logo_base64 = load_image_as_base64(logo_path)
        if logo_base64:
            logo_html = (
                f'<img src="data:image/png;base64,{logo_base64}" '
                f'style="height:20px; vertical-align:middle; margin-top: 0;" alt="Inventify Logo"/>'
            )
            st.markdown(
                f"""
                <div style="text-align: center; margin-top: 3em; padding-top: 1rem; border-top: 1px solid #e9ecef;">
                    <p>Resume ATS Optimizer Pro © 2025 | Powered by {logo_html}</p>
                </div>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.warning(f"Logo file not found at path: {logo_path}")

if __name__ == "__main__":
    main()
